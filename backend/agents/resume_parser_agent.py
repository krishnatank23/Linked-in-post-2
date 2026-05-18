import os
import re
import traceback
from typing import Any
import fitz
from docx import Document
from langchain_openai import ChatOpenAI
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from env_config import load_backend_env
from path_resolver import resolve_resume_path
from agents.json_utils import parse_llm_json_content

import asyncio
import json
from agents.llm_guard import guarded_llm_ainvoke
load_backend_env()


def _get_deepseek_api_key() -> str | None:
    """Support the DeepSeek API key env var."""
    return os.getenv("DEEPSEEK_API_KEY")


def _get_groq_api_key() -> str | None:
    """Support the Groq API key env var."""
    return os.getenv("GROQ_API_KEY")

RESUME_PARSER_PROMPT = """You are an expert resume and LinkedIn profile analyzer. 
Your job is to extract ALL structured information from the given resume/profile text.

Extract the following information in a structured JSON format:

{{
    "personal_info": {{
        "full_name": "...",
        "email": "...",
        "phone": "...",
        "location": "...",
        "linkedin_url": "...",
        "portfolio_url": "..."
    }},
    "professional_summary": "A brief professional summary from the resume. IF there is no explicit summary in the resume, YOU MUST read the rest of the resume and write a concise, highly accurate 2-3 sentence summary about the user.",
    "experience": [
        {{
            "company": "...",
            "role": "...",
            "duration": "...",
            "description": "...",
            "key_achievements": ["..."]
        }}
    ],
    "education": [
        {{
            "institution": "...",
            "degree": "...",
            "field_of_study": "...",
            "year": "..."
        }}
    ],
    "skills": {{
        "technical_skills": ["..."],
        "soft_skills": ["..."],
        "tools_and_technologies": ["..."],
        "languages": ["..."]
    }},
    "certifications": ["..."],
    "projects": [
        {{
            "name": "...",
            "description": "...",
            "technologies_used": ["..."]
        }}
    ],
    "achievements_and_awards": ["..."],
    "interests": ["..."],
    "total_years_of_experience": "...",
    "current_role": "...",
    "industry": "...",
    "expertise_areas": ["..."]
}}

If any field is not found in the resume, use null or an empty array. 
EXCEPTION: For "professional_summary", if it is missing, you MUST synthesize an accurate one yourself based on their experience and skills. Do NOT return null for the summary.
Be thorough and extract every possible detail.

Resume Text:
{resume_text}

Return ONLY the JSON object, no markdown fences, no extra text.
"""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    return text.strip()


def extract_resume_text(file_path: str) -> str:
    """Extract text from one resume/profile file based on extension."""
    resolved_path = resolve_resume_path(file_path)
    ext = os.path.splitext(resolved_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(resolved_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(resolved_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def extract_resume_text_bundle(file_paths: str | list[str]) -> tuple[str, int]:
    """Extract and combine text from one or multiple uploaded profile documents."""
    if isinstance(file_paths, list):
        normalized_paths = [str(path).strip() for path in file_paths if str(path or "").strip()]
    else:
        normalized_paths = [str(file_paths).strip()] if str(file_paths or "").strip() else []

    if not normalized_paths:
        raise ValueError("No resume/profile document path provided")

    chunks: list[str] = []
    for idx, path in enumerate(normalized_paths, start=1):
        extracted = extract_resume_text(path)
        if extracted:
            chunks.append(f"\n\n=== DOCUMENT {idx} START ===\n{extracted}\n=== DOCUMENT {idx} END ===")

    combined_text = "\n".join(chunks).strip()
    return combined_text, len(normalized_paths)


def _extract_first_email(text: str) -> str | None:
    """Extract the first syntactically valid email from resume text."""
    email_pattern = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
    match = email_pattern.search(text or "")
    return match.group(0).strip() if match else None


def _extract_first_url_by_hint(text: str, hint: str) -> str | None:
    """Extract first URL containing a hint token (e.g. linkedin.com)."""
    url_pattern = re.compile(r"(?:https?://)?(?:www\.)?[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s]*)?")
    for match in url_pattern.finditer(text or ""):
        candidate = match.group(0).strip().rstrip(".,;)")
        if hint.lower() in candidate.lower():
            if not candidate.startswith(("http://", "https://")):
                candidate = f"https://{candidate}"
            return candidate
    return None


def _sanitize_personal_info(parsed_data: dict[str, Any], resume_text: str) -> None:
    """Normalize critical personal fields to avoid malformed outputs from OCR/LLM noise."""
    if not isinstance(parsed_data, dict):
        return

    personal_info = parsed_data.get("personal_info")
    if not isinstance(personal_info, dict):
        personal_info = {}
        parsed_data["personal_info"] = personal_info

    # Email: prefer strict regex extraction from source text.
    extracted_email = _extract_first_email(resume_text)
    if extracted_email:
        personal_info["email"] = extracted_email
    else:
        raw_email = str(personal_info.get("email") or "").strip()
        llm_email_match = _extract_first_email(raw_email)
        personal_info["email"] = llm_email_match if llm_email_match else None

    # LinkedIn URL: if model output is malformed, recover from source text.
    raw_linkedin = str(personal_info.get("linkedin_url") or "").strip()
    if "linkedin.com" not in raw_linkedin.lower():
        extracted_linkedin = _extract_first_url_by_hint(resume_text, "linkedin.com")
        personal_info["linkedin_url"] = extracted_linkedin if extracted_linkedin else (raw_linkedin or None)

    # Portfolio URL: keep best effort, but normalize obvious malformed values.
    raw_portfolio = str(personal_info.get("portfolio_url") or "").strip()
    if raw_portfolio and raw_portfolio.lower().startswith("http"):
        personal_info["portfolio_url"] = raw_portfolio.rstrip(".,;)")


# Local parsing functions removed in favor of agents.json_utils.parse_llm_json_content


async def run_resume_parser(file_path: str | list[str]) -> dict[str, Any]:
    """
    Agent 1: Parse resume and extract structured data using DeepSeek LLM.
    Returns a dict with status, output, and optional error.
    """
    try:
        # Step 1: Extract raw text from one or more uploaded profile documents.
        print(f"[DEBUG] Resume Parser: Extracting text from {file_path}...")
        resume_text, source_document_count = extract_resume_text_bundle(file_path)
        print(f"[DEBUG] Resume Parser: Extracted {len(resume_text)} chars from {source_document_count} document(s)")

        if not resume_text or len(resume_text.strip()) < 50:
            return {
                "status": "error",
                "output": None,
                "error": "Could not extract sufficient text from the resume. Please upload a valid PDF or DOCX file.",
            }

        # Step 2: Build the LLM chain with optional fallback
        # ds_key = _get_deepseek_api_key()
        groq_key = _get_groq_api_key()
        prompt = ChatPromptTemplate.from_template(RESUME_PARSER_PROMPT)
        
        chains = []
        # if ds_key:
        #     primary_llm = ChatOpenAI(
        #         model=os.getenv("RESUME_PARSER_MODEL", "deepseek-v4-flash"),
        #         temperature=0.1,
        #         api_key=ds_key,
        #         base_url="https://api.deepseek.com",
        #     )
        #     chains.append(prompt | primary_llm)
        
        if groq_key:
            fallback_llm = ChatGroq(
                model=os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
                temperature=0.1,
                groq_api_key=groq_key,
            )
            chains.append(prompt | fallback_llm)

        if not chains:
            return {
                "status": "error",
                "output": None,
                "error": "No AI API keys (DeepSeek or Groq) found in environment.",
            }

        chain = chains[0].with_fallbacks(chains[1:]) if len(chains) > 1 else chains[0]

        try:
            print(f"[DEBUG] Resume Parser: Invoking LLM to parse resume (timeout: 45s)...")
            response = await guarded_llm_ainvoke(
                chain,
                {"resume_text": resume_text},
                timeout_seconds=45,
            )
            print(f"[DEBUG] Resume Parser: LLM response received")
        except asyncio.TimeoutError:
            return {
                "status": "error",
                "output": None,
                "error": "Resume parsing timed out after 45 seconds. The LLM service may be slow.",
            }

        content = response.content.strip()
        if not content:
            return {
                "status": "error",
                "output": None,
                "error": "The AI model returned an empty response during resume parsing. Please try again.",
            }

        parsed_data = parse_llm_json_content(content)
        if not isinstance(parsed_data, dict):
            raise json.JSONDecodeError("Resume parser output was not a JSON object", str(parsed_data), 0)

        print(f"[DEBUG] Resume Parser: Successfully parsed resume JSON")
        _sanitize_personal_info(parsed_data, resume_text)

        return {
            "status": "success",
            "output": {
                "raw_text_length": len(resume_text),
                "source_document_count": source_document_count,
                "parsed_profile": parsed_data,
            },
            "error": None,
        }

    except json.JSONDecodeError as e:
        return {
            "status": "error",
            "output": {"raw_response": response.content if 'response' in locals() else "N/A"},
            "error": f"Failed to parse LLM response as JSON: {str(e)}",
        }
    except Exception as e:
        err = str(e)
        if "rate limit" in err.lower() or "429" in err:
            return {
                "status": "error",
                "output": None,
                "error": "DeepSeek API limit reached. Please try again later.",
            }
        return {
            "status": "error",
            "output": None,
            "error": f"Resume parser failed: {str(e)}\n{traceback.format_exc()}",
        }
