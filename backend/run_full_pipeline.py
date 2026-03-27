import asyncio
import uuid
from pathlib import Path

import aiohttp
from docx import Document


API_BASE = "http://localhost:8000/api"
API_ROOT = API_BASE.rsplit("/api", 1)[0]


def create_test_resume_docx(file_path: Path) -> None:
    doc = Document()
    doc.add_heading("Krishna Tank", 0)
    doc.add_paragraph("Email: krishna@example.com")
    doc.add_paragraph("Phone: +91-9999999999")
    doc.add_paragraph("Location: India")
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Software engineer with 5 years of experience building backend APIs and AI-enabled products.")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Backend Developer, Example Tech (2021-Present)")
    doc.add_paragraph("Built FastAPI services, integrated LLM APIs, improved reliability and observability.")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.Tech in Computer Science")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, SQLAlchemy, JavaScript, Docker, LangChain")
    doc.save(file_path)


async def check_backend_available(session: aiohttp.ClientSession) -> bool:
    try:
        async with session.get(f"{API_ROOT}/openapi.json") as resp:
            if resp.status == 200:
                return True
            print(f"   backend preflight failed: unexpected status {resp.status}")
            return False
    except Exception as exc:
        print("   backend preflight failed: could not reach API server")
        print("   start backend first: python backend/main.py")
        print(f"   details: {exc}")
        return False


async def main() -> None:
    unique = str(uuid.uuid4())[:8]
    email = f"pipeline_{unique}@test.com"
    username = f"pipeline_{unique}"
    password = "test12345"

    resume_path = Path("tmp_resume.docx")
    create_test_resume_docx(resume_path)

    timeout = aiohttp.ClientTimeout(total=130)
    try:
        async with aiohttp.ClientSession(timeout=timeout) as session:
            print("0) Checking backend availability...")
            if not await check_backend_available(session):
                return

            print("1) Registering test user...")
            with resume_path.open("rb") as f:
                form = aiohttp.FormData()
                form.add_field("email", email)
                form.add_field("username", username)
                form.add_field("password", password)
                form.add_field("resume", f, filename="resume.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                async with session.post(f"{API_BASE}/register", data=form) as resp:
                    reg_body = await resp.text()
                    print(f"   register status={resp.status}")
                    if resp.status != 200:
                        print(f"   register body={reg_body}")
                        return
                    reg_json = await resp.json()
                    user_id = reg_json["user_id"]

            print("2) Logging in...")
            async with session.post(
                f"{API_BASE}/login",
                json={"email": email, "password": password},
                headers={"Content-Type": "application/json"},
            ) as resp:
                login_body = await resp.text()
                print(f"   login status={resp.status}")
                if resp.status != 200:
                    print(f"   login body={login_body}")
                    return
                login_json = await resp.json()
                token = login_json["access_token"]

            print("3) Running full pipeline...")
            async with session.post(
                f"{API_BASE}/pipeline/run",
                json={"user_id": user_id},
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {token}"},
            ) as resp:
                run_body = await resp.text()
                print(f"   pipeline run status={resp.status}")
                if resp.status != 200:
                    print(f"   pipeline run body={run_body}")
                else:
                    run_json = await resp.json()
                    print(f"   pipeline message={run_json.get('message')}")
                    for item in run_json.get("results", []):
                        print(f"   - {item.get('agent_name')}: {item.get('status')}")
                        if item.get("error"):
                            print(f"     error: {item.get('error')}")

            print("4) Fetching pipeline results...")
            async with session.get(
                f"{API_BASE}/pipeline/results/{user_id}",
                headers={"Authorization": f"Bearer {token}"},
            ) as resp:
                results_body = await resp.text()
                print(f"   results status={resp.status}")
                if resp.status != 200:
                    print(f"   results body={results_body}")
                else:
                    results_json = await resp.json()
                    print(f"   stored results count={len(results_json.get('results', []))}")
    finally:
        if resume_path.exists():
            resume_path.unlink()


if __name__ == "__main__":
    asyncio.run(main())
